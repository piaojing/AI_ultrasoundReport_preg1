"""
Second Trimester Ultrasound Report Generator
============================================
GUI  : Tkinter
AI   : Google Gemini 2.5 Flash (free tier)
Fallback AI : OpenRouter NVIDIA Nemotron 3 Ultra (free tier) when Gemini rate limits are hit
Output: Professional .docx  named  PatientName-AgeGender.docx

Requirements:
    pip install google-generativeai python-docx requests

Free Gemini API key: https://aistudio.google.com/apikey
"""

import tkinter as tk
from tkinter import ttk, messagebox, filedialog
import threading
import os
import re
import time
from datetime import date
from dotenv import load_dotenv, set_key
import requests

# ── third-party ──────────────────────────────────────────────────────────────
try:
    import google.generativeai as genai
except ImportError:
    raise SystemExit("Missing package. Run:  pip install google-generativeai")

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    raise SystemExit("Missing package. Run:  pip install python-docx")

try:
    from PIL import Image, ImageTk
except ImportError:
    Image = None
    ImageTk = None

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DOTENV_PATH = os.path.join(BASE_DIR, ".env")
load_dotenv(DOTENV_PATH)
RECORD_DIR = os.path.join(BASE_DIR, "record")

DEFAULT_SAVE_DIR = os.getenv("SAVE_DIR") or os.path.expanduser("~\\Documents")


def save_dir_to_env(path: str):
    try:
        set_key(DOTENV_PATH, "SAVE_DIR", path)
    except Exception:
        with open(DOTENV_PATH, "a", encoding="utf-8") as env_file:
            env_file.write(f"\nSAVE_DIR={path}\n")

# ─────────────────────────────────────────────────────────────────────────────
#  CONFIGURATION
# ─────────────────────────────────────────────────────────────────────────────
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")   # set in .env file  # https://aistudio.google.com/apikey
GEMINI_MODEL   = "gemini-2.5-flash"
OPENROUTER_API_KEY = os.getenv("OPENROUTER_API_KEY")  # set in .env file if you want fallback to OpenRouter
OPENROUTER_MODEL = os.getenv("OPENROUTER_MODEL") or "openai/gpt-oss-120b:free"
OPENROUTER_API_BASE = os.getenv("OPENROUTER_API_BASE") or "https://openrouter.ai/api/v1"
if "api.openrouter.ai" in OPENROUTER_API_BASE:
    OPENROUTER_API_BASE = "https://openrouter.ai/api/v1"
OPENROUTER_API_BASE = OPENROUTER_API_BASE.rstrip("/")


class GeminiRateLimitError(RuntimeError):
    """Raised when Gemini free tier limit is reached and fallback is needed."""


# Section headings used to detect headings in AI output
SECTION_HEADINGS = {
    "PATIENT INFORMATION", "CLINICAL INDICATION", "EXAMINATION",
    "TECHNIQUE", "FINDINGS", "IMPRESSION", "RECOMMENDATION",
    "CONCLUSION", "REPORT"
}

with open("reffer.txt") as f:
    SECOND_TRIMESTER_REFERENCES = f.read()

BLUE   = RGBColor(0x18, 0x5F, 0xA5)
DKBLUE = RGBColor(0x0C, 0x44, 0x7C)
GRAY   = RGBColor(0x55, 0x55, 0x55)
LGRAY  = RGBColor(0xAA, 0xAA, 0xAA)

# ─────────────────────────────────────────────────────────────────────────────
#  GEMINI CALL  (with retry on 429)
# ─────────────────────────────────────────────────────────────────────────────

def call_gemini(prompt: str, status_cb=None) -> str:
    genai.configure(api_key=GEMINI_API_KEY)
    model = genai.GenerativeModel(GEMINI_MODEL)

    try:
        if status_cb:
            status_cb("Calling Gemini AI …")
        response = model.generate_content(prompt)
        return response.text.strip()

    except Exception as exc:
        msg = str(exc)
        if "429" in msg or "ResourceExhausted" in msg or "quota" in msg.lower():
            raise GeminiRateLimitError(msg) from exc
        raise


def call_openrouter(prompt: str, status_cb=None) -> str:
    if status_cb:
        status_cb(f"Calling OpenRouter {OPENROUTER_MODEL}…")

    if not OPENROUTER_API_KEY:
        raise RuntimeError(
            "OpenRouter API key is not configured. Set OPENROUTER_API_KEY in .env to enable fallback."
        )

    url = f"{OPENROUTER_API_BASE}/chat/completions"
    headers = {
        "Authorization": f"Bearer {OPENROUTER_API_KEY}",
        "Content-Type": "application/json",
    }
    payload = {
        "model": OPENROUTER_MODEL,
        "messages": [{"role": "user", "content": prompt}],
        "temperature": 0.0,
        "max_tokens": 1200,
    }

    try:
        response = requests.post(url, headers=headers, json=payload, timeout=120)
        response.raise_for_status()
        data = response.json()

        if not data or "choices" not in data or not data["choices"]:
            raise RuntimeError("OpenRouter response did not return any choices.")

        choice = data["choices"][0]
        if isinstance(choice.get("message"), dict):
            content = choice["message"].get("content")
            if isinstance(content, str):
                return content.strip()

        if isinstance(choice.get("text"), str):
            return choice["text"].strip()

        raise RuntimeError(
            "Unable to parse OpenRouter completion response. "
            "The completion choice did not include valid text content."
        )

    except requests.exceptions.RequestException as exc:
        raise RuntimeError(
            "OpenRouter request failed. Check OPENROUTER_API_BASE, OPENROUTER_API_KEY, and network connectivity. "
            f"Tried: {url}. Last error: {exc}"
        ) from exc


# ─────────────────────────────────────────────────────────────────────────────
#  PROMPT BUILDER
# ─────────────────────────────────────────────────────────────────────────────

def build_prompt(data: dict) -> str:
    """Convert GUI form data into a detailed radiologist prompt."""

    p = data   # shorthand

    # Build measurements block only for non-empty fields
    def row(label, val, unit=""):
        if val and val.strip() and val.strip() not in ("-", "N/A", ""):
            return f"  {label}: {val.strip()} {unit}".rstrip()
        return None

    meas_lines = list(filter(None, [
        row("Gestational Age by LMP",          p.get("ga_lmp"),       "weeks"),
        row("Biparietal Diameter (BPD)",        p.get("bpd"),          "mm"),
        row("Head Circumference (HC)",          p.get("hc"),           "mm"),
        row("Abdominal Circumference (AC)",     p.get("ac"),           "mm"),
        row("Femur Length (FL)",                p.get("fl"),           "mm"),
        row("Pregnancy Presentation",           p.get("presentation")),
        row("Fetal Heart Rate",                 p.get("fhr"),          "bpm"),
        row("Lateral ventricular size",         p.get("ventricular_size")),
        row("Cerebellar diameter",              p.get("cerebellar_diameter")),
        row("CM size",                          p.get("cm_size")),
        row("Nose length",                      p.get("nose_length")),
        row("Facial / jaw / extremities",       p.get("facial_jaw_extremities")),
        row("Skull bones",                      p.get("skull_bones")),
        row("Spine / skeletal system",          p.get("spine_skeletal")),
        row("Placenta position",                p.get("placenta_position")),
        row("Placenta thickness",               p.get("placenta_thickness"), "mm"),
        row("Placenta maturity",                p.get("placenta_maturity")),
        row("Placenta edges",                   p.get("placenta_edges")),
        row("Amniotic Fluid",                   p.get("amniotic")),
        row("Umbilical Cord",                   p.get("umbilical_cord")),
        row("Umbilical artery Doppler RI",      p.get("umbilical_artery_ri")),
        row("Uterus",                           p.get("uterus")),
        row("Adnexa / Ovaries",                 p.get("adnexa")),
        row("Other parts",                      p.get("other_parts")),
        row("Additional Findings",              p.get("additional")),
    ]))

    measurements = "\n".join(meas_lines) if meas_lines else "  (No specific measurements provided)"

    return f"""You are a senior radiologist specialising in obstetric ultrasound.
Write a complete, formal SECOND TRIMESTER OBSTETRIC ULTRASOUND REPORT.

Use these exact section headings (all caps, no colon):
PATIENT INFORMATION
CLINICAL INDICATION
TECHNIQUE
FINDINGS
IMPRESSION

Patient details:
  Name            : {p.get('name', 'Unknown')}
  Age             : {p.get('age', '')} years
  Gender          : {p.get('gender', 'Female')}
  Exam Date       : {p.get('exam_date', date.today().strftime('%d/%m/%Y'))}
  Referring MD    : {p.get('referring_md', '')}
  G/P             : G{p.get('gravida','?')} P{p.get('para','?')}
  LMP             : {p.get('lmp', '')}
  Indication      : {p.get('indication', 'Routine second trimester scan')}
  Scan Type       : {p.get('scan_type', 'Transabdominal')}

Ultrasound measurements and observations:
{measurements}

Writing rules:
- Expand raw data into complete, professional radiological sentences.
- Create a formal radiology report tone appropriate for obstetric ultrasound.
- FINDINGS section must address each observed parameter individually and clearly.
- IMPRESSION must summarise fetal viability, key findings, and any relevant second trimester observations.
- If any field is blank or missing, omit the specific detail gracefully without inventing data.
- Write in plain text only — no markdown, no asterisks, no bullet characters.
- Be medically precise but concise.
- Do NOT add any text outside the five sections above.
"""


def translate_report_to_vietnamese(report_text: str) -> str:
    return (
        "Translate the following second trimester obstetric ultrasound report into Vietnamese. "
        "Preserve the exact section headings in uppercase and keep the same report structure. "
        "Do not add any sections or commentary beyond the translated report text.\n\n"
        f"{report_text}"
    )


def _add_report_text(doc: Document, report_text: str):
    for line in report_text.splitlines():
        stripped = line.strip()

        if not stripped:
            sp = doc.add_paragraph()
            sp.paragraph_format.space_after = Pt(2)
            continue

        heading_key = stripped.upper().rstrip(":")
        if heading_key in SECTION_HEADINGS:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(stripped.rstrip(":").upper())
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = BLUE
            run.font.name = "Arial"
            _add_rule(p, "CCCCCC", 4)

        elif stripped.endswith(":") and len(stripped) < 45:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(stripped)
            run.bold = True
            run.font.size = Pt(10)
            run.font.name = "Arial"

        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(stripped)
            run.font.size = Pt(10)
            run.font.name = "Arial"


# ─────────────────────────────────────────────────────────────────────────────
#  DOCX BUILDER
# ─────────────────────────────────────────────────────────────────────────────

def _add_rule(paragraph, color_hex="185FA5", size=6):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)


def build_docx(english_text: str, vietnamese_text: str, patient: dict) -> Document:
    doc = Document()

    # --- Page setup (A4, 2 cm margins) ---
    sec = doc.sections[0]
    sec.top_margin    = Cm(2)
    sec.bottom_margin = Cm(2)
    sec.left_margin   = Cm(2.5)
    sec.right_margin  = Cm(2.5)

    # ── Hospital / title block ─────────────────────────────────────────────
    hosp = doc.add_paragraph()
    hosp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = hosp.add_run("DEPARTMENT OF RADIOLOGY")
    r.bold = True; r.font.size = Pt(13); r.font.color.rgb = DKBLUE; r.font.name = "Arial"

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub.add_run("Second Trimester Obstetric Ultrasound Report")
    r2.font.size = Pt(10); r2.font.color.rgb = GRAY; r2.font.name = "Arial"
    _add_rule(sub, "185FA5", 8)
    sub.paragraph_format.space_after = Pt(10)

    # ── English report ────────────────────────────────────────────────────
    _add_report_text(doc, english_text)

    # ── Vietnamese report ─────────────────────────────────────────────────
    doc.add_page_break()
    vn_title = doc.add_paragraph()
    vn_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_vn = vn_title.add_run("BÁO CÁO TIẾNG VIỆT")
    run_vn.bold = True
    run_vn.font.size = Pt(11)
    run_vn.font.color.rgb = DKBLUE
    run_vn.font.name = "Arial"
    vn_title.paragraph_format.space_after = Pt(8)

    _add_report_text(doc, vietnamese_text)

    # ── Signature block ────────────────────────────────────────────────────
    doc.add_paragraph()
    sig = doc.add_paragraph()
    _add_rule(sig, "CCCCCC", 4)
    sig.paragraph_format.space_before = Pt(16)
    run = sig.add_run("Reported by:  ________________________________")
    run.font.size = Pt(9); run.font.color.rgb = GRAY; run.font.name = "Arial"

    dt = doc.add_paragraph()
    run2 = dt.add_run(
        f"Date: {patient.get('exam_date', date.today().strftime('%d/%m/%Y'))}"
    )
    run2.font.size = Pt(9); run2.font.color.rgb = GRAY; run2.font.name = "Arial"

    return doc

# ─────────────────────────────────────────────────────────────────────────────
#  FILENAME HELPER
# ─────────────────────────────────────────────────────────────────────────────

def safe_filename(patient: dict) -> str:
    name   = re.sub(r"[^\w\s-]", "", patient.get("name", "Patient")).strip()
    name   = re.sub(r"\s+", "_", name)
    age    = patient.get("age", "")
    gender = patient.get("gender", "F")[0].upper()
    return f"{name}-{age}{gender}.docx"

# ─────────────────────────────────────────────────────────────────────────────
#  TKINTER GUI
# ─────────────────────────────────────────────────────────────────────────────

class App(tk.Tk):
    FIELD_PAD = {"padx": 6, "pady": 3}

    def __init__(self):
        super().__init__()
        self.title("Second Trimester US Report Generator")
        self.resizable(True, True)
        self.configure(bg="#F4F6FA")

        # ----- top bar -------------------------------------------------------
        topbar = tk.Frame(self, bg="#185FA5")
        topbar.pack(fill="x")
        tk.Label(
            topbar,
            text="  🔬  Second Trimester Ultrasound Report  |  AI-Powered",
            bg="#185FA5", fg="white",
            font=("Arial", 12, "bold"), anchor="w", pady=8
        ).pack(side="left")

        # ----- scrollable canvas ---------------------------------------------
        container = tk.Frame(self, bg="#F4F6FA")
        container.pack(fill="both", expand=True, side="left")

        canvas = tk.Canvas(container, bg="#F4F6FA", highlightthickness=0)
        scrollbar = ttk.Scrollbar(container, orient="vertical", command=canvas.yview)
        self.scroll_frame = tk.Frame(canvas, bg="#F4F6FA")

        ref_panel = tk.Frame(container, bg="#F4F6FA", width=520)
        ref_panel.pack(side="right", fill="y", padx=(0, 2))
        ref_panel.pack_propagate(False)

        tk.Label(
            ref_panel, text="Second Trimester Reference",
            bg="#185FA5", fg="white",
            font=("Arial", 10, "bold"), anchor="center", pady=8
        ).pack(fill="x")

        if Image and ImageTk:
            try:
                img_path = os.path.join(BASE_DIR, "ref.jpg")
                img = Image.open(img_path)
                max_width = 500
                if img.width > max_width:
                    target_height = round(img.height * max_width / img.width)
                    img = img.resize((max_width, target_height), Image.LANCZOS)
                self._ref_image = ImageTk.PhotoImage(img)

                ref_image_canvas = tk.Canvas(
                    ref_panel, bg="white", bd=1, relief="solid",
                    highlightthickness=0, xscrollincrement=1, yscrollincrement=1
                )
                ref_image_canvas.create_image(0, 0, anchor="nw", image=self._ref_image)
                ref_image_canvas.configure(
                    scrollregion=(0, 0, self._ref_image.width(), self._ref_image.height())
                )
                ref_image_canvas.pack(fill="both", expand=True, padx=4, pady=(4, 0))

                ref_hscroll = ttk.Scrollbar(ref_panel, orient="horizontal", command=ref_image_canvas.xview)
                ref_hscroll.pack(fill="x", side="bottom", padx=4, pady=(0, 4))
                ref_image_canvas.configure(xscrollcommand=ref_hscroll.set)

                ref_vscroll = ttk.Scrollbar(ref_panel, orient="vertical", command=ref_image_canvas.yview)
                ref_vscroll.pack(fill="y", side="right", padx=(0, 4), pady=4)
                ref_image_canvas.configure(yscrollcommand=ref_vscroll.set)
            except Exception as exc:
                tk.Label(
                    ref_panel,
                    text=f"Unable to load ref.jpg: {exc}",
                    bg="white", fg="red",
                    wraplength=500, justify="left"
                ).pack(fill="both", expand=True, padx=4, pady=4)
        else:
            tk.Label(
                ref_panel,
                text="Pillow is required to display ref.jpg. Install via pip install pillow.",
                bg="white", fg="red",
                wraplength=500, justify="left"
            ).pack(fill="both", expand=True, padx=4, pady=4)

        self.scroll_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )
        canvas.create_window((0, 0), window=self.scroll_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)

        scrollbar.pack(side="right", fill="y")
        ref_panel.pack(side="right", fill="y", padx=(0, 2))
        canvas.pack(side="left", fill="both", expand=True)

        # Mouse-wheel scroll
        self.bind_all("<MouseWheel>",
                      lambda e: canvas.yview_scroll(int(-1*(e.delta/120)), "units"))

        # ----- build form sections -------------------------------------------
        self.vars = {}
        self._build_form()

        # ----- status bar + button -------------------------------------------
        bottom = tk.Frame(self, bg="#E8ECF4", pady=8)
        bottom.pack(fill="x", side="bottom")

        self.status_var = tk.StringVar(value="Ready. Fill in the form and click Generate.")
        tk.Label(
            bottom, textvariable=self.status_var,
            bg="#E8ECF4", fg="#444", font=("Arial", 9), anchor="w"
        ).pack(side="top", padx=10)

        self.gen_btn = tk.Button(
            bottom,
            text="  ▶  Generate Report",
            command=self._on_generate,
            bg="#185FA5", fg="white",
            font=("Arial", 11, "bold"),
            relief="flat", padx=14, pady=6,
            cursor="hand2",
            activebackground="#0C447C", activeforeground="white"
        )
        self.gen_btn.pack(side="right", padx=10)

        self.geometry("760x800")
        self.minsize(680, 600)
        self.state("zoomed")

    # ── form builder helpers ──────────────────────────────────────────────

    def _section(self, title: str) -> tk.Frame:
        outer = tk.Frame(self.scroll_frame, bg="#F4F6FA")
        outer.pack(fill="x", padx=10, pady=(10, 0))

        header = tk.Frame(outer, bg="#185FA5")
        header.pack(fill="x")
        tk.Label(
            header, text=f"  {title}",
            bg="#185FA5", fg="white",
            font=("Arial", 9, "bold"), anchor="w", pady=4
        ).pack(side="left")

        body = tk.LabelFrame(outer, bg="white", relief="groove", bd=1)
        body.pack(fill="x", pady=(0, 0))
        return body

    def _row(self, parent, label: str, key: str, width=28,
             choices=None, default=""):
        """Single label + entry (or combobox) row."""
        row = tk.Frame(parent, bg="white")
        row.pack(fill="x", padx=8, pady=2)
        tk.Label(
            row, text=label, bg="white",
            font=("Arial", 9), width=28, anchor="w"
        ).pack(side="left")

        var = tk.StringVar(value=default)
        self.vars[key] = var

        if choices:
            w = ttk.Combobox(row, textvariable=var, values=choices,
                             width=width, font=("Arial", 9))
            w.set(default or choices[0])
        else:
            w = tk.Entry(row, textvariable=var, width=width+2,
                         font=("Arial", 9), relief="groove", bd=1)
        w.pack(side="left", padx=4)
        return w

    def _two_col(self, parent, items):
        """Two entries side by side: items = [(label, key, width, choices), ...]"""
        row = tk.Frame(parent, bg="white")
        row.pack(fill="x", padx=8, pady=2)
        for label, key, width, choices in items:
            tk.Label(row, text=label, bg="white",
                     font=("Arial", 9), width=22, anchor="w").pack(side="left")
            var = tk.StringVar()
            self.vars[key] = var
            if choices:
                w = ttk.Combobox(row, textvariable=var, values=choices,
                                 width=width, font=("Arial", 9))
                w.set(choices[0])
            else:
                w = tk.Entry(row, textvariable=var, width=width,
                             font=("Arial", 9), relief="groove", bd=1)
            w.pack(side="left", padx=(4, 12))

    # ── form sections ─────────────────────────────────────────────────────

    def _build_form(self):
        # # 1. API KEY
        # api_sec = self._section("⚙  API Configuration")
        # self._row(api_sec, "Gemini API Key *", "api_key",
        #           width=42, default=GEMINI_API_KEY)
        # tk.Label(api_sec,
        #          text="  Get free key → https://aistudio.google.com/apikey",
        #          bg="white", fg="#185FA5", font=("Arial", 8)).pack(anchor="w", padx=8)

        # 2. PATIENT INFO
        pat = self._section("👤  Patient Information")
        self._two_col(pat, [
            ("Patient Name *",  "name",     22, None),
            ("Age (years) *",   "age",      6,  None),
        ])
        self._two_col(pat, [
            ("Gender",          "gender",   10, ["Female", "Male"]),
            ("Exam Date",       "exam_date",12, None),
        ])
        self._two_col(pat, [
            ("Gravida",         "gravida",  5,  [str(i) for i in range(0, 15)]),
            ("Para",            "para",     5,  [str(i) for i in range(0, 15)]),
        ])
        self._row(pat, "LMP (dd/mm/yyyy)",       "lmp",         width=14)
        self._row(pat, "Referring Physician",     "referring_md",width=28)
        self._row(pat, "Clinical Indication",     "indication",  width=42,
                  default="Routine second trimester scan")

        # 3. SCAN PARAMETERS
        scan = self._section("📡  Scan Parameters")
        self._row(scan, "Scan Type", "scan_type", choices=[
            "Transabdominal", "Transvaginal", "Transabdominal + Transvaginal"
        ])

        # 4. BIOMETRY & FETAL ANATOMY
        bio = self._section("📐  Biometry & Fetal Anatomy")
        self._two_col(bio, [
            ("BPD (mm)", "bpd", 7, None),
            ("HC (mm)",  "hc",  7, None),
        ])
        self._two_col(bio, [
            ("AC (mm)", "ac", 7, None),
            ("FL (mm)", "fl", 7, None),
        ])
        self._row(bio, "Lateral ventricular size", "ventricular_size",
                  width=40, default="Normal (<10 mm)")
        self._two_col(bio, [
            ("Cerebellar diameter", "cerebellar_diameter", 20, None),
            ("CM size",             "cm_size",               20, None),
        ])
        self._two_col(bio, [
            ("Nose length", "nose_length", 20, None),
            ("Skull bones", "skull_bones", 20, None),
        ])
        self._row(bio, "Facial / jaw / extremities", "facial_jaw_extremities",
                  width=40, default="No abnormal images currently displayed")
        self._row(bio, "Spine / skeletal system", "spine_skeletal",
                  width=40, default="No abnormal images currently displayed")

        # 5. FETAL WELLBEING
        fw = self._section("❤  Fetal Wellbeing")
        self._two_col(fw, [
            ("Presentation", "presentation", 18,
             ["Not fixed", "Cephalic", "Breech", "Transverse", "Oblique", ""]),
            ("Fetal Heart Rate (bpm)", "fhr", 7, None),
        ])
        self._row(fw, "Umbilical artery Doppler RI", "umbilical_artery_ri",
                  width=40)

        # 6. PLACENTA
        pl = self._section("🩺  Placenta")
        self._row(pl, "Placenta position", "placenta_position",
                  width=40, default="Posterior, attached to the back wall of the uterus")
        self._row(pl, "Placenta thickness (mm)", "placenta_thickness",
                  width=40)
        self._two_col(pl, [
            ("Placenta maturity", "placenta_maturity", 20, None),
            ("Placenta edges",    "placenta_edges",    20, None),
        ])

        # 7. UTERUS, OVARIES & AMNIOTIC FLUID
        pu = self._section("🔬  Uterus, Ovaries & Amniotic Fluid")
        self._row(pu, "Uterus / Myometrium", "uterus",
                  width=40, default="Uniform structure, no mass")
        self._row(pu, "Adnexa / Ovaries", "adnexa",
                  width=40, default="No abnormal images on ovaries")
        self._row(pu, "Amniotic Fluid", "amniotic",
                  width=40, default="Clear fluid, normal amount")

        # 8. UMBILICAL CORD & OTHER PARTS
        uc = self._section("🧬  Umbilical Cord & Other Parts")
        self._row(uc, "Umbilical Cord", "umbilical_cord",
                  width=40, default="02 arteries, 01 vein")
        self._row(uc, "Other parts", "other_parts",
                  width=40, default="Normal")

        # 9. ADDITIONAL
        add = self._section("📝  Additional Findings / Comments")
        self._row(add, "Additional findings", "additional", width=52)

        # 9. OUTPUT
        out = self._section("💾  Output")
        self._row(out, "Save to folder", "save_dir", width=40,
                  default=DEFAULT_SAVE_DIR)
        tk.Button(
            out, text="Browse…",
            command=self._browse_dir,
            font=("Arial", 8), relief="groove", bg="#E8ECF4"
        ).pack(side="left", padx=(0, 8), pady=4)

        # Set today's date as default
        self.vars["exam_date"].set(date.today().strftime("%d/%m/%Y"))

    # ── events ────────────────────────────────────────────────────────────

    def _browse_dir(self):
        d = filedialog.askdirectory(title="Select output folder")
        if d:
            self.vars["save_dir"].set(d)
            save_dir_to_env(d)

    def _on_generate(self):
        api_var = self.vars.get("api_key")
        api_key = api_var.get().strip() if api_var else ""
        api_key = api_key or GEMINI_API_KEY or ""
        if not api_key:
            messagebox.showwarning(
                "API Key",
                "Please enter your Gemini API key or set GEMINI_API_KEY in .env."
            )
            return

        self.gen_btn.config(state="disabled", text="  ⏳  Working…")
        threading.Thread(target=self._generate_thread,
                         args=(api_key,), daemon=True).start()

    def _set_status(self, msg: str):
        self.status_var.set(msg)
        self.update_idletasks()

    def _show_info_dialog(self, title: str, message: str):
        event = threading.Event()
        dialog_result = {"shown": False}

        def show():
            messagebox.showinfo(title, message)
            dialog_result["shown"] = True
            event.set()

        self.after(0, show)
        event.wait()
        return dialog_result["shown"]

    def _generate_thread(self, api_key: str):
        try:
            # Collect all field values
            patient = {k: v.get() for k, v in self.vars.items()}
            patient["api_key"] = api_key      # keep for later

            # Override global API key with user-entered one
            global GEMINI_API_KEY
            GEMINI_API_KEY = api_key

            prompt = build_prompt(patient)
            try:
                report_text = call_gemini(prompt, status_cb=self._set_status)
            except GeminiRateLimitError:
                if OPENROUTER_API_KEY:
                    self._set_status("Gemini free limit reached. Switching to OpenRouter…")
                    self._show_info_dialog(
                        "Gemini Limit Reached",
                        "Gemini free quota was exceeded (429). Click OK to switch to OpenRouter NVIDIA Nemotron 3 Ultra fallback."
                    )
                    report_text = call_openrouter(prompt, status_cb=self._set_status)
                else:
                    raise RuntimeError(
                        "Gemini free quota exceeded. Set OPENROUTER_API_KEY in .env to enable fallback."
                    )

            vn_prompt = translate_report_to_vietnamese(report_text)
            try:
                report_text_vn = call_gemini(vn_prompt, status_cb=self._set_status)
            except GeminiRateLimitError:
                if OPENROUTER_API_KEY:
                    self._set_status("Gemini free limit reached during translation. Switching to OpenRouter…")
                    self._show_info_dialog(
                        "Gemini Limit Reached",
                        "Gemini free quota was exceeded during translation. Click OK to switch to OpenRouter NVIDIA Nemotron 3 Ultra fallback."
                    )
                    report_text_vn = call_openrouter(vn_prompt, status_cb=self._set_status)
                else:
                    raise RuntimeError(
                        "Gemini free quota exceeded. Set OPENROUTER_API_KEY in .env to enable fallback."
                    )

            self._set_status("Building Word document…")
            docx_doc = build_docx(report_text, report_text_vn, patient)

            fname    = safe_filename(patient)
            save_dir = patient.get("save_dir", "").strip() or DEFAULT_SAVE_DIR
            os.makedirs(save_dir, exist_ok=True)
            filepath = os.path.join(save_dir, fname)
            docx_doc.save(filepath)

            try:
                os.startfile(filepath)
                self._set_status(f"✅  Saved and opened: {filepath}")
            except Exception as err:
                self._set_status(f"Saved, but could not open Word: {err}")

            self.after(500, self.destroy)

        except Exception as exc:
            self._set_status(f"❌  Error: {exc}")
            messagebox.showerror("Error", str(exc))
            self.gen_btn.config(state="normal", text="  ▶  Generate Report")


# ─────────────────────────────────────────────────────────────────────────────
#  ENTRY POINT
# ─────────────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    app = App()
    app.mainloop()
