"""
First Trimester Ultrasound Report Generator
============================================
GUI  : Tkinter
AI   : Google Gemini 2.5 Flash (free tier)
Output: Professional .docx  named  PatientName-AgeGender.docx

Requirements:
    pip install google-generativeai python-docx

Free Gemini API key: https://aistudio.google.com/apikey
"""

import tkinter as tk
from tkinter import ttk, messagebox, filedialog
import threading
import os
import re
import time
from datetime import date
from dotenv import load_dotenv, set_key

# ── third-party ──────────────────────────────────────────────────────────────
try:
    import google.generativeai as genai
except ImportError:
    raise SystemExit("Missing package. Run:  pip install google-generativeai")

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    raise SystemExit("Missing package. Run:  pip install python-docx")

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DOTENV_PATH = os.path.join(BASE_DIR, ".env")
load_dotenv(DOTENV_PATH)
RECORD_DIR = os.path.join(BASE_DIR, "record")

DEFAULT_SAVE_DIR = os.getenv("SAVE_DIR") or os.path.expanduser("~\\Documents")


def save_dir_to_env(path: str):
    try:
        set_key(DOTENV_PATH, "SAVE_DIR", path)
    except Exception:
        with open(DOTENV_PATH, "a", encoding="utf-8") as env_file:
            env_file.write(f"\nSAVE_DIR={path}\n")

# ─────────────────────────────────────────────────────────────────────────────
#  CONFIGURATION
# ─────────────────────────────────────────────────────────────────────────────
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")   # set in .env file  # https://aistudio.google.com/apikey
GEMINI_MODEL   = "gemini-2.5-flash"
MAX_RETRIES    = 5
RETRY_WAIT     = 65          # seconds between retries on 429

# Section headings used to detect headings in AI output
SECTION_HEADINGS = {
    "PATIENT INFORMATION", "CLINICAL INDICATION", "EXAMINATION",
    "TECHNIQUE", "FINDINGS", "IMPRESSION", "RECOMMENDATION",
    "CONCLUSION", "REPORT"
}

FIRST_TRIMESTER_REFERENCES = (
    "First Trimester Reference Ranges:\n"
    "  - GA by LMP: 11 to 13+6 weeks\n"
    "  - CRL: 45 to 84 mm\n"
    "  - BPD: approx. 18 to 31 mm\n"
    "  - HC: approx. 70 to 110 mm\n"
    "  - AC: approx. 60 to 95 mm\n"
    "  - FL: approx. 10 to 30 mm\n"
    "  - NT: normal <3.0 mm (typical 1.0 to 2.5 mm)\n"
    "  - FHR: 110 to 160 bpm\n"
    "  - Yolk sac: 2 to 6 mm\n"
    "  - Cervical length: >25 mm\n"
    "  - Placenta: normal location, no previa\n\n"
    "Use these as general first trimester references.\n"
    "If a value is unavailable, leave it blank rather than guessing."
)

BLUE   = RGBColor(0x18, 0x5F, 0xA5)
DKBLUE = RGBColor(0x0C, 0x44, 0x7C)
GRAY   = RGBColor(0x55, 0x55, 0x55)
LGRAY  = RGBColor(0xAA, 0xAA, 0xAA)

# ─────────────────────────────────────────────────────────────────────────────
#  GEMINI CALL  (with retry on 429)
# ─────────────────────────────────────────────────────────────────────────────

def call_gemini(prompt: str, status_cb=None) -> str:
    genai.configure(api_key=GEMINI_API_KEY)
    model = genai.GenerativeModel(GEMINI_MODEL)

    for attempt in range(1, MAX_RETRIES + 1):
        try:
            if status_cb:
                status_cb(f"Calling Gemini AI … (attempt {attempt})")
            response = model.generate_content(prompt)
            return response.text.strip()

        except Exception as exc:
            msg = str(exc)
            if "429" in msg or "ResourceExhausted" in msg or "quota" in msg.lower():
                # Try to parse retry delay from the message
                wait = RETRY_WAIT
                m = re.search(r"retry[_ ]in[^\d]*(\d+)", msg, re.I)
                if m:
                    wait = int(m.group(1)) + 3

                if attempt < MAX_RETRIES:
                    for remaining in range(wait, 0, -1):
                        if status_cb:
                            status_cb(
                                f"Rate limit hit — retrying in {remaining}s "
                                f"(attempt {attempt}/{MAX_RETRIES})"
                            )
                        time.sleep(1)
                    continue
                else:
                    raise RuntimeError(
                        f"Gemini rate limit: all {MAX_RETRIES} attempts failed.\n"
                        "Wait a minute and try again."
                    ) from exc
            else:
                raise

# ─────────────────────────────────────────────────────────────────────────────
#  PROMPT BUILDER
# ─────────────────────────────────────────────────────────────────────────────

def build_prompt(data: dict) -> str:
    """Convert GUI form data into a detailed radiologist prompt."""

    p = data   # shorthand

    # Build measurements block only for non-empty fields
    def row(label, val, unit=""):
        if val and val.strip() and val.strip() not in ("-", "N/A", ""):
            return f"  {label}: {val.strip()} {unit}".rstrip()
        return None

    meas_lines = list(filter(None, [
        row("Gestational Age by LMP",          p.get("ga_lmp"),       "weeks"),
        row("Crown-Rump Length (CRL)",          p.get("crl"),          "mm"),
        row("Gestational Age by CRL",           p.get("ga_crl"),       "weeks+days"),
        row("Biparietal Diameter (BPD)",        p.get("bpd"),          "mm"),
        row("Head Circumference (HC)",          p.get("hc"),           "mm"),
        row("Abdominal Circumference (AC)",     p.get("ac"),           "mm"),
        row("Femur Length (FL)",                p.get("fl"),           "mm"),
        row("Nuchal Translucency (NT)",         p.get("nt"),           "mm"),
        row("Nasal Bone",                       p.get("nasal_bone")),
        row("Heart Rate",                       p.get("fhr"),          "bpm"),
        row("Cardiac Activity",                 p.get("cardiac")),
        row("Fetal Movement",                   p.get("movement")),
        row("Fetal Number",                     p.get("fetal_num")),
        row("Chorionicity (if multiple)",       p.get("chorionicity")),
        row("Placenta Location",                p.get("placenta_loc")),
        row("Placenta Appearance",              p.get("placenta_app")),
        row("Myometrium",                       p.get("myometrium")),
        row("Cervical Length",                  p.get("cx_length"),    "mm"),
        row("Amniotic Fluid",                   p.get("amniotic")),
        row("Adnexa / Ovaries",                 p.get("adnexa")),
        row("Yolk Sac",                         p.get("yolk_sac")),
        row("Gestational Sac Diameter (MSD)",   p.get("msd"),          "mm"),
        row("Uterus",                           p.get("uterus")),
        row("Additional Findings",              p.get("additional")),
    ]))

    measurements = "\n".join(meas_lines) if meas_lines else "  (No specific measurements provided)"

    return f"""You are a senior radiologist specialising in obstetric ultrasound.
Write a complete, formal FIRST TRIMESTER OBSTETRIC ULTRASOUND REPORT.

Use these exact section headings (all caps, no colon):
PATIENT INFORMATION
CLINICAL INDICATION
TECHNIQUE
FINDINGS
IMPRESSION

Patient details:
  Name            : {p.get('name', 'Unknown')}
  Age             : {p.get('age', '')} years
  Gender          : {p.get('gender', 'Female')}
  Exam Date       : {p.get('exam_date', date.today().strftime('%d/%m/%Y'))}
  Referring MD    : {p.get('referring_md', '')}
  G/P             : G{p.get('gravida','?')} P{p.get('para','?')}
  LMP             : {p.get('lmp', '')}
  Indication      : {p.get('indication', 'Routine first trimester scan')}
  Scan Type       : {p.get('scan_type', 'Transabdominal')}

Ultrasound measurements and observations:
{measurements}

Writing rules:
- Expand raw data into complete, professional radiological sentences.
- Create a formal radiology report tone appropriate for obstetric ultrasound.
- FINDINGS section must address each measured parameter individually and clearly.
- IMPRESSION must summarise gestational age, viability, and any key findings;
  end with an estimated due date (EDD) calculated from CRL-based GA or GS based GA or AC based GA if available,
  otherwise from LMP.
- If any field is blank or missing, omit the specific detail gracefully without
  inventing data.
- Write in plain text only — no markdown, no asterisks, no bullet characters.
- Be medically precise but concise.
- Do NOT add any text outside the five sections above.
"""


def translate_report_to_vietnamese(report_text: str) -> str:
    return (
        "Translate the following first trimester obstetric ultrasound report into Vietnamese. "
        "Preserve the exact section headings in uppercase and keep the same report structure. "
        "Do not add any sections or commentary beyond the translated report text.\n\n"
        f"{report_text}"
    )


def _add_report_text(doc: Document, report_text: str):
    for line in report_text.splitlines():
        stripped = line.strip()

        if not stripped:
            sp = doc.add_paragraph()
            sp.paragraph_format.space_after = Pt(2)
            continue

        heading_key = stripped.upper().rstrip(":")
        if heading_key in SECTION_HEADINGS:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(stripped.rstrip(":").upper())
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = BLUE
            run.font.name = "Arial"
            _add_rule(p, "CCCCCC", 4)

        elif stripped.endswith(":") and len(stripped) < 45:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(stripped)
            run.bold = True
            run.font.size = Pt(10)
            run.font.name = "Arial"

        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(stripped)
            run.font.size = Pt(10)
            run.font.name = "Arial"


# ─────────────────────────────────────────────────────────────────────────────
#  DOCX BUILDER
# ─────────────────────────────────────────────────────────────────────────────

def _add_rule(paragraph, color_hex="185FA5", size=6):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)


def build_docx(english_text: str, vietnamese_text: str, patient: dict) -> Document:
    doc = Document()

    # --- Page setup (A4, 2 cm margins) ---
    sec = doc.sections[0]
    sec.top_margin    = Cm(2)
    sec.bottom_margin = Cm(2)
    sec.left_margin   = Cm(2.5)
    sec.right_margin  = Cm(2.5)

    # ── Hospital / title block ─────────────────────────────────────────────
    hosp = doc.add_paragraph()
    hosp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = hosp.add_run("DEPARTMENT OF RADIOLOGY")
    r.bold = True; r.font.size = Pt(13); r.font.color.rgb = DKBLUE; r.font.name = "Arial"

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub.add_run("First Trimester Obstetric Ultrasound Report")
    r2.font.size = Pt(10); r2.font.color.rgb = GRAY; r2.font.name = "Arial"
    _add_rule(sub, "185FA5", 8)
    sub.paragraph_format.space_after = Pt(10)

    # ── English report ────────────────────────────────────────────────────
    _add_report_text(doc, english_text)

    # ── Vietnamese report ─────────────────────────────────────────────────
    doc.add_page_break()
    vn_title = doc.add_paragraph()
    vn_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_vn = vn_title.add_run("BÁO CÁO TIẾNG VIỆT")
    run_vn.bold = True
    run_vn.font.size = Pt(11)
    run_vn.font.color.rgb = DKBLUE
    run_vn.font.name = "Arial"
    vn_title.paragraph_format.space_after = Pt(8)

    _add_report_text(doc, vietnamese_text)

    # ── Signature block ────────────────────────────────────────────────────
    doc.add_paragraph()
    sig = doc.add_paragraph()
    _add_rule(sig, "CCCCCC", 4)
    sig.paragraph_format.space_before = Pt(16)
    run = sig.add_run("Reported by:  ________________________________")
    run.font.size = Pt(9); run.font.color.rgb = GRAY; run.font.name = "Arial"

    dt = doc.add_paragraph()
    run2 = dt.add_run(
        f"Date: {patient.get('exam_date', date.today().strftime('%d/%m/%Y'))}"
    )
    run2.font.size = Pt(9); run2.font.color.rgb = GRAY; run2.font.name = "Arial"

    return doc

# ─────────────────────────────────────────────────────────────────────────────
#  FILENAME HELPER
# ─────────────────────────────────────────────────────────────────────────────

def safe_filename(patient: dict) -> str:
    name   = re.sub(r"[^\w\s-]", "", patient.get("name", "Patient")).strip()
    name   = re.sub(r"\s+", "_", name)
    age    = patient.get("age", "")
    gender = patient.get("gender", "F")[0].upper()
    return f"{name}-{age}{gender}.docx"

# ─────────────────────────────────────────────────────────────────────────────
#  TKINTER GUI
# ─────────────────────────────────────────────────────────────────────────────

class App(tk.Tk):
    FIELD_PAD = {"padx": 6, "pady": 3}

    def __init__(self):
        super().__init__()
        self.title("First Trimester US Report Generator")
        self.resizable(True, True)
        self.configure(bg="#F4F6FA")

        # ----- top bar -------------------------------------------------------
        topbar = tk.Frame(self, bg="#185FA5")
        topbar.pack(fill="x")
        tk.Label(
            topbar,
            text="  🔬  First Trimester Ultrasound Report  |  AI-Powered",
            bg="#185FA5", fg="white",
            font=("Arial", 12, "bold"), anchor="w", pady=8
        ).pack(side="left")

        # ----- scrollable canvas ---------------------------------------------
        container = tk.Frame(self, bg="#F4F6FA")
        container.pack(fill="both", expand=True)

        canvas = tk.Canvas(container, bg="#F4F6FA", highlightthickness=0)
        scrollbar = ttk.Scrollbar(container, orient="vertical", command=canvas.yview)
        self.scroll_frame = tk.Frame(canvas, bg="#F4F6FA")

        ref_panel = tk.Frame(container, bg="#F4F6FA", width=300)
        ref_panel.pack(side="right", fill="y")
        ref_panel.pack_propagate(False)

        tk.Label(
            ref_panel, text="First Trimester Reference",
            bg="#185FA5", fg="white",
            font=("Arial", 10, "bold"), anchor="center", pady=8
        ).pack(fill="x")

        ref_text = tk.Text(
            ref_panel, bg="white", fg="#333",
            font=("Arial", 9), bd=1, relief="solid",
            wrap="word", state="normal", padx=8, pady=8
        )
        ref_text.insert("1.0", FIRST_TRIMESTER_REFERENCES)
        ref_text.configure(state="disabled")
        ref_text.pack(fill="both", expand=True, padx=8, pady=8)

        self.scroll_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )
        canvas.create_window((0, 0), window=self.scroll_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)

        scrollbar.pack(side="right", fill="y")
        canvas.pack(side="left", fill="both", expand=True)

        # Mouse-wheel scroll
        self.bind_all("<MouseWheel>",
                      lambda e: canvas.yview_scroll(int(-1*(e.delta/120)), "units"))

        # ----- build form sections -------------------------------------------
        self.vars = {}
        self._build_form()

        # ----- status bar + button -------------------------------------------
        bottom = tk.Frame(self, bg="#E8ECF4", pady=8)
        bottom.pack(fill="x", side="bottom")

        self.status_var = tk.StringVar(value="Ready. Fill in the form and click Generate.")
        tk.Label(
            bottom, textvariable=self.status_var,
            bg="#E8ECF4", fg="#444", font=("Arial", 9), anchor="w"
        ).pack(side="left", padx=10)

        self.gen_btn = tk.Button(
            bottom,
            text="  ▶  Generate Report",
            command=self._on_generate,
            bg="#185FA5", fg="white",
            font=("Arial", 11, "bold"),
            relief="flat", padx=14, pady=6,
            cursor="hand2",
            activebackground="#0C447C", activeforeground="white"
        )
        self.gen_btn.pack(side="right", padx=10)

        self.geometry("760x800")
        self.minsize(680, 600)
        self.state("zoomed")

    # ── form builder helpers ──────────────────────────────────────────────

    def _section(self, title: str) -> tk.Frame:
        outer = tk.Frame(self.scroll_frame, bg="#F4F6FA")
        outer.pack(fill="x", padx=10, pady=(10, 0))

        header = tk.Frame(outer, bg="#185FA5")
        header.pack(fill="x")
        tk.Label(
            header, text=f"  {title}",
            bg="#185FA5", fg="white",
            font=("Arial", 9, "bold"), anchor="w", pady=4
        ).pack(side="left")

        body = tk.LabelFrame(outer, bg="white", relief="groove", bd=1)
        body.pack(fill="x", pady=(0, 0))
        return body

    def _row(self, parent, label: str, key: str, width=28,
             choices=None, default=""):
        """Single label + entry (or combobox) row."""
        row = tk.Frame(parent, bg="white")
        row.pack(fill="x", padx=8, pady=2)
        tk.Label(
            row, text=label, bg="white",
            font=("Arial", 9), width=28, anchor="w"
        ).pack(side="left")

        var = tk.StringVar(value=default)
        self.vars[key] = var

        if choices:
            w = ttk.Combobox(row, textvariable=var, values=choices,
                             width=width, font=("Arial", 9))
            w.set(default or choices[0])
        else:
            w = tk.Entry(row, textvariable=var, width=width+2,
                         font=("Arial", 9), relief="groove", bd=1)
        w.pack(side="left", padx=4)
        return w

    def _two_col(self, parent, items):
        """Two entries side by side: items = [(label, key, width, choices), ...]"""
        row = tk.Frame(parent, bg="white")
        row.pack(fill="x", padx=8, pady=2)
        for label, key, width, choices in items:
            tk.Label(row, text=label, bg="white",
                     font=("Arial", 9), width=22, anchor="w").pack(side="left")
            var = tk.StringVar()
            self.vars[key] = var
            if choices:
                w = ttk.Combobox(row, textvariable=var, values=choices,
                                 width=width, font=("Arial", 9))
                w.set(choices[0])
            else:
                w = tk.Entry(row, textvariable=var, width=width,
                             font=("Arial", 9), relief="groove", bd=1)
            w.pack(side="left", padx=(4, 12))

    # ── form sections ─────────────────────────────────────────────────────

    def _build_form(self):
        # # 1. API KEY
        # api_sec = self._section("⚙  API Configuration")
        # self._row(api_sec, "Gemini API Key *", "api_key",
        #           width=42, default=GEMINI_API_KEY)
        # tk.Label(api_sec,
        #          text="  Get free key → https://aistudio.google.com/apikey",
        #          bg="white", fg="#185FA5", font=("Arial", 8)).pack(anchor="w", padx=8)

        # 2. PATIENT INFO
        pat = self._section("👤  Patient Information")
        self._two_col(pat, [
            ("Patient Name *",  "name",     22, None),
            ("Age (years) *",   "age",      6,  None),
        ])
        self._two_col(pat, [
            ("Gender",          "gender",   10, ["Female", "Male"]),
            ("Exam Date",       "exam_date",12, None),
        ])
        self._two_col(pat, [
            ("Gravida",         "gravida",  5,  [str(i) for i in range(0, 15)]),
            ("Para",            "para",     5,  [str(i) for i in range(0, 15)]),
        ])
        self._row(pat, "LMP (dd/mm/yyyy)",       "lmp",         width=14)
        self._row(pat, "Referring Physician",     "referring_md",width=28)
        self._row(pat, "Clinical Indication",     "indication",  width=42,
                  default="Routine first trimester scan")

        # 3. SCAN PARAMETERS
        scan = self._section("📡  Scan Parameters")
        self._row(scan, "Scan Type", "scan_type", choices=[
            "Transabdominal", "Transvaginal", "Transabdominal + Transvaginal"
        ])

        # 4. GESTATIONAL AGE & BIOMETRY
        bio = self._section("📐  Biometry & Gestational Age")
        self._two_col(bio, [
            ("MSD — gestational sac (mm)", "msd", 7,  None),
            ("CRL (mm)",                "crl",    7,  None),
        ])
        self._two_col(bio, [
            ("BPD (mm)",                "bpd",    7,  None),
            ("HC (mm)",                 "hc",     7,  None),
        ])
        self._two_col(bio, [
            ("AC (mm)",                 "ac",     7,  None),
            ("FL (mm)",                 "fl",     7,  None),
        ])

        # 5. NT & MARKERS
        nt = self._section("🔎  Nuchal Translucency & First-Trimester Markers")
        self._two_col(nt, [
            ("Nuchal Translucency (mm)", "nt",         7, None),
            ("Nasal Bone",               "nasal_bone", 16,
             ["Present", "Absent", "Not clearly visualised", ""]),
        ])

        # 6. FETAL WELLBEING
        fw = self._section("❤  Fetal Wellbeing")
        self._two_col(fw, [
            ("Cardiac Activity",    "cardiac",
             12, ["Present", "Absent", "Regular", "Irregular", ""]),
            ("Fetal Heart Rate (bpm)", "fhr", 7, None),
        ])
        self._two_col(fw, [
            ("Fetal Movement",      "movement",
             12, ["Present", "Absent", "Active", "Not seen", ""]),
            ("Yolk Sac",            "yolk_sac",
             16, ["Visible, normal", "Not visualised", "Enlarged", ""]),
        ])
        self._two_col(fw, [
            ("Fetal Number",        "fetal_num",
             10, ["Singleton", "Twins", "Triplets", ""]),
            ("Chorionicity",        "chorionicity",
             20, ["N/A", "Dichorionic-Diamniotic", "Monochorionic-Diamniotic",
                  "Monochorionic-Monoamniotic", ""]),
        ])

        # 7. PLACENTA & UTERUS
        pu = self._section("🔬  Uterus & Cervix")
        self._row(pu, "Cervical Length (mm)",  "cervical_length",
                  width=40,
                  default="Normal size and echogenicity, no fibroids")
        self._row(pu, "Uterus / Myometrium",  "uterus",
                  width=40,
                  default="Normal size and echogenicity, no fibroids")
        self._row(pu, "Adnexa / Ovaries",     "adnexa",
                  width=40, default="Unremarkable bilateral ovaries")
  
        # 8. ADDITIONAL
        add = self._section("📝  Additional Findings / Comments")
        self._row(add, "Additional findings", "additional", width=52)

        # 9. OUTPUT
        out = self._section("💾  Output")
        self._row(out, "Save to folder", "save_dir", width=40,
                  default=DEFAULT_SAVE_DIR)
        tk.Button(
            out, text="Browse…",
            command=self._browse_dir,
            font=("Arial", 8), relief="groove", bg="#E8ECF4"
        ).pack(side="left", padx=(0, 8), pady=4)

        # Set today's date as default
        self.vars["exam_date"].set(date.today().strftime("%d/%m/%Y"))

    # ── events ────────────────────────────────────────────────────────────

    def _browse_dir(self):
        d = filedialog.askdirectory(title="Select output folder")
        if d:
            self.vars["save_dir"].set(d)
            save_dir_to_env(d)

    def _on_generate(self):
        api_var = self.vars.get("api_key")
        api_key = api_var.get().strip() if api_var else ""
        api_key = api_key or GEMINI_API_KEY or ""
        if not api_key:
            messagebox.showwarning(
                "API Key",
                "Please enter your Gemini API key or set GEMINI_API_KEY in .env."
            )
            return

        self.gen_btn.config(state="disabled", text="  ⏳  Working…")
        threading.Thread(target=self._generate_thread,
                         args=(api_key,), daemon=True).start()

    def _set_status(self, msg: str):
        self.status_var.set(msg)
        self.update_idletasks()

    def _generate_thread(self, api_key: str):
        try:
            # Collect all field values
            patient = {k: v.get() for k, v in self.vars.items()}
            patient["api_key"] = api_key      # keep for later

            # Override global API key with user-entered one
            global GEMINI_API_KEY
            GEMINI_API_KEY = api_key

            prompt      = build_prompt(patient)
            report_text = call_gemini(prompt, status_cb=self._set_status)
            vn_prompt   = translate_report_to_vietnamese(report_text)
            report_text_vn = call_gemini(vn_prompt, status_cb=self._set_status)

            self._set_status("Building Word document…")
            docx_doc    = build_docx(report_text, report_text_vn, patient)

            fname    = safe_filename(patient)
            save_dir = patient.get("save_dir", "").strip() or DEFAULT_SAVE_DIR
            os.makedirs(save_dir, exist_ok=True)
            filepath = os.path.join(save_dir, fname)
            docx_doc.save(filepath)

            try:
                os.startfile(filepath)
                self._set_status(f"✅  Saved and opened: {filepath}")
            except Exception as err:
                self._set_status(f"Saved, but could not open Word: {err}")

            self.after(500, self.destroy)

        except Exception as exc:
            self._set_status(f"❌  Error: {exc}")
            messagebox.showerror("Error", str(exc))
            self.gen_btn.config(state="normal", text="  ▶  Generate Report")


# ─────────────────────────────────────────────────────────────────────────────
#  ENTRY POINT
# ─────────────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    app = App()
    app.mainloop()
